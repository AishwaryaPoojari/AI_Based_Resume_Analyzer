import streamlit as st
import io
import os
import re as _re
from validators import is_valid_email, is_valid_phone, is_valid_url, is_valid_linkedin_url, is_valid_github_url


def load_css():
    with open("styles/style.css", encoding="utf-8") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)


load_css()

# ── Auth Guard ─────────────────────────────────────────────
if not st.session_state.get("logged_in"):
    st.warning("⚠️ Please log in first.")
    if st.button("Go to Login"):
        st.switch_page("pages/Login.py")
    st.stop()


def generate_resume_docx(data):
    """Generate resume using python-docx — no Node.js needed."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def hex_to_rgb(h):
        h = h.lstrip("#")
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    def add_hr(doc, color_hex="2E75B6"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_section_heading(doc, text, color_hex):
        p   = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold      = True
        run.font.size = Pt(12)
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
        run.font.name = "Arial"
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def add_body(doc, text, bold_part=None, bold_text=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        if bold_part and bold_text:
            run1 = p.add_run(bold_part + ": ")
            run1.bold      = True
            run1.font.size = Pt(10)
            run1.font.name = "Arial"
            run2 = p.add_run(bold_text)
            run2.font.size = Pt(10)
            run2.font.name = "Arial"
        else:
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = "Arial"
        return p

    def add_bullet(doc, text):
        p   = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Arial"
        p.paragraph_format.space_after = Pt(2)
        return p

    color = data.get("color", "2E75B6")
    r, g, b = hex_to_rgb(color)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # ── Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(data["name"].upper())
    name_run.bold           = True
    name_run.font.size      = Pt(22)
    name_run.font.name      = "Arial"
    name_run.font.color.rgb = RGBColor(r, g, b)
    name_p.paragraph_format.space_after = Pt(2)

    # ── Contact
    contact_parts = [data.get("email",""), data.get("phone",""), data.get("linkedin",""), data.get("github","")]
    contact_line  = "  |  ".join(p for p in contact_parts if p.strip())
    if contact_line:
        c_p = doc.add_paragraph()
        c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_run = c_p.add_run(contact_line)
        c_run.font.size      = Pt(9)
        c_run.font.name      = "Arial"
        c_run.font.color.rgb = RGBColor(85, 85, 85)
        c_p.paragraph_format.space_after = Pt(2)

    # ── Role line
    role_line = "  |  ".join(p for p in [data.get("occupation",""), data.get("department","")] if p.strip())
    if role_line:
        r_p = doc.add_paragraph()
        r_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_run = r_p.add_run(role_line)
        r_run.italic         = True
        r_run.font.size      = Pt(10)
        r_run.font.name      = "Arial"
        r_run.font.color.rgb = RGBColor(119, 119, 119)
        r_p.paragraph_format.space_after = Pt(4)

    add_hr(doc, color)

    # ── Summary
    if data.get("summary"):
        add_section_heading(doc, "Professional Summary", color)
        add_body(doc, data["summary"])
        add_hr(doc, color)

    # ── Education
    edu_list = [e for e in data.get("education", []) if e.get("degree","").strip()]
    if edu_list:
        add_section_heading(doc, "Education", color)
        for e in edu_list:
            p = doc.add_paragraph()
            run = p.add_run(e["degree"])
            run.bold      = True
            run.font.size = Pt(11)
            run.font.name = "Arial"
            p.paragraph_format.space_after = Pt(2)
            line = "  |  ".join(x for x in [e.get("college",""), e.get("year",""), e.get("grade","")] if x.strip())
            if line:
                add_body(doc, line)
        add_hr(doc, color)

    # ── Skills
    if data.get("skills"):
        add_section_heading(doc, "Technical Skills", color)
        add_body(doc, "", bold_part="Technical", bold_text=data["skills"])
        if data.get("softskills"):
            add_body(doc, "", bold_part="Soft Skills", bold_text=data["softskills"])
        add_hr(doc, color)

    # ── Experience
    if data.get("experience"):
        add_section_heading(doc, "Experience Level", color)
        add_body(doc, data["experience"])
        add_hr(doc, color)

    # ── Projects
    projects = [p for p in data.get("projects", []) if p.get("name","").strip()]
    if projects:
        add_section_heading(doc, "Projects", color)
        for proj in projects:
            p = doc.add_paragraph()
            run = p.add_run(proj["name"])
            run.bold      = True
            run.font.size = Pt(11)
            run.font.name = "Arial"
            p.paragraph_format.space_after = Pt(2)
            if proj.get("description"):
                for line in proj["description"].split("\n"):
                    if line.strip():
                        add_bullet(doc, line.strip())
            if proj.get("tech"):
                add_body(doc, "", bold_part="Tech Stack", bold_text=proj["tech"])
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        add_hr(doc, color)

    # ── Certifications
    if data.get("certifications"):
        add_section_heading(doc, "Certifications & Achievements", color)
        for line in data["certifications"].split("\n"):
            if line.strip():
                add_bullet(doc, line.strip())
        add_hr(doc, color)

    # ── Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("Generated by AI Resume Analyzer")
    footer_run.italic         = True
    footer_run.font.size      = Pt(8)
    footer_run.font.color.rgb = RGBColor(170, 170, 170)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ── Header ─────────────────────────────────────────────────
c1, c2, c3 = st.columns([7, 2, 2])
with c1:
    st.markdown("## 🛠️ Resume Generator")
    st.markdown("<p style='color:#6B7280'>Fill in your details and download a professional resume instantly.</p>",
                unsafe_allow_html=True)
with c2:
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("📊 Dashboard", use_container_width=True):
        st.switch_page("pages/Dashboard.py")
with c3:
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("🚪 Logout", use_container_width=True):
        st.session_state.clear()
        st.switch_page("pages/Landing.py")

st.markdown("<hr style='border-color:rgba(139,92,246,0.14)'>", unsafe_allow_html=True)

st.markdown("""
<div style="background:#EDE9FF;border:1px solid rgba(139,92,246,0.3);border-radius:10px;
            padding:0.8rem 1.2rem;margin-bottom:1.5rem;">
    <b style="color:#7C3AED;">&#128161; How it works:</b>
    <span style="color:#6B7280;font-size:0.87rem;">
        Your basic info is pre-filled from your account.
        Add your skills, projects and education — then click Generate!
    </span>
</div>
""", unsafe_allow_html=True)

# ── Color Theme ────────────────────────────────────────────
st.markdown("### 🎨 Choose Resume Color Theme")
t1, t2, t3, t4 = st.columns(4)
with t1:
    if st.button("🔵 Blue",   use_container_width=True): st.session_state.resume_color = "2E75B6"
with t2:
    if st.button("🟢 Green",  use_container_width=True): st.session_state.resume_color = "00695c"
with t3:
    if st.button("⚫ Dark",   use_container_width=True): st.session_state.resume_color = "1a1a2e"
with t4:
    if st.button("🔴 Maroon", use_container_width=True): st.session_state.resume_color = "7b0000"

chosen_color = st.session_state.get("resume_color", "2E75B6")
color_names  = {"2E75B6": "Blue", "00695c": "Green", "1a1a2e": "Dark", "7b0000": "Maroon"}
st.markdown(f"<p style='color:#6B7280;font-size:0.82rem;'>Selected: <b style='color:#1E1245;'>{color_names.get(chosen_color,'Blue')}</b></p>",
            unsafe_allow_html=True)

st.markdown("<hr style='border-color:rgba(139,92,246,0.14)'>", unsafe_allow_html=True)

# ── Personal Details ───────────────────────────────────────
st.markdown("### 👤 Personal Details")

row1 = st.columns(2)
with row1[0]:
    name = st.text_input("Full Name *", value=st.session_state.get("user_name", ""))
    if name.strip() and any(ch.isdigit() for ch in name):
        st.markdown(
            "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ Name should not contain numbers.</p>",
            unsafe_allow_html=True
        )
with row1[1]:
    phone = st.text_input("Phone Number", placeholder="9000000000", max_chars=10)
    if phone.strip():
        if not phone.strip().isdigit():
            st.markdown(
                "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ Phone number can only contain digits (no letters).</p>",
                unsafe_allow_html=True
            )
        elif len(phone.strip()) != 10:
            st.markdown(
                "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ Phone number must be exactly 10 digits.</p>",
                unsafe_allow_html=True
            )

row2 = st.columns(2)
with row2[0]:
    email = st.text_input("Email *", value=st.session_state.get("user_email", ""))
    if email.strip() and not is_valid_email(email.strip()):
        st.markdown(
            "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ Please enter a valid email address.</p>",
            unsafe_allow_html=True
        )
with row2[1]:
    linkedin = st.text_input("LinkedIn URL", placeholder="linkedin.com/in/yourname")
    if linkedin.strip() and not is_valid_linkedin_url(linkedin.strip()):
        st.markdown(
            "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ Please enter a valid LinkedIn URL (e.g. linkedin.com/in/yourname).</p>",
            unsafe_allow_html=True
        )

github = st.text_input("GitHub URL (optional)", placeholder="github.com/yourname")
if github.strip() and not is_valid_github_url(github.strip()):
    st.markdown(
        "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ Please enter a valid GitHub URL (e.g. github.com/yourname).</p>",
        unsafe_allow_html=True
    )

st.markdown("<hr style='border-color:rgba(139,92,246,0.14)'>", unsafe_allow_html=True)

# ── Professional Details ───────────────────────────────────
st.markdown("### 💼 Professional Details")

DEPARTMENT_OPTIONS = [
    "Select Department", "Computer Science", "Information Technology",
    "Electronics & Communication", "Mechanical Engineering", "Civil Engineering",
    "Electrical Engineering", "Business Administration", "Commerce",
    "Arts & Humanities", "Other"
]
EXPERIENCE_OPTIONS = [
    "Select Experience", "Fresher (0 years)", "0–1 years",
    "1–2 years", "2–5 years", "5–10 years", "10+ years"
]
OCCUPATION_OPTIONS = [
    "Select Occupation", "Student", "Recent Graduate", "Software Developer",
    "Data Analyst", "Data Scientist", "Web Developer", "DevOps Engineer",
    "Business Analyst", "HR Professional", "Project Manager", "Other"
]


def _preselect(options, saved_value):
    """Index of the account's saved value in the dropdown, or 0 (the
    'Select ...' placeholder) if it isn't one of the listed options."""
    return options.index(saved_value) if saved_value in options else 0


pr1, pr2 = st.columns(2)
with pr1:
    department = st.selectbox(
        "Department *", DEPARTMENT_OPTIONS,
        index=_preselect(DEPARTMENT_OPTIONS, st.session_state.get("user_department", ""))
    )
with pr2:
    experience = st.selectbox(
        "Experience Level *", EXPERIENCE_OPTIONS,
        index=_preselect(EXPERIENCE_OPTIONS, st.session_state.get("user_experience", ""))
    )

occupation = st.selectbox(
    "Current Occupation *", OCCUPATION_OPTIONS,
    index=_preselect(OCCUPATION_OPTIONS, st.session_state.get("user_occupation", ""))
)

summary = st.text_area("Professional Summary *",
                        placeholder="Write 2-3 sentences about yourself and your career goal.",
                        height=100)

_summary_sentences = [s for s in _re.split(r"[.!?]+", summary.strip()) if s.strip()]
if summary.strip() and len(_summary_sentences) < 2:
    st.markdown(
        "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ Please write more than one sentence about yourself.</p>",
        unsafe_allow_html=True
    )

st.markdown("<hr style='border-color:rgba(139,92,246,0.14)'>", unsafe_allow_html=True)

# ── Education ──────────────────────────────────────────────
st.markdown("### 🎓 Education")
education = []
for idx in range(2):
    label = "Primary" if idx == 0 else "Secondary (optional)"
    with st.expander(f"📚 {label} Education", expanded=(idx == 0)):
        row_a = st.columns(2)
        with row_a[0]:
            degree = st.text_input("Degree / Course", placeholder="B.Tech Computer Science", key=f"deg_{idx}")
            if degree.strip() and any(ch.isdigit() for ch in degree):
                st.markdown(
                    "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ Degree should contain letters only, no numbers.</p>",
                    unsafe_allow_html=True
                )
        with row_a[1]:
            year = st.text_input("Year", placeholder="2020 - 2025", key=f"yr_{idx}")
            if year.strip() and not _re.match(r"^\d{4}\s*-\s*\d{4}$", year.strip()):
                st.markdown(
                    "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ Use the format 2020 - 2025.</p>",
                    unsafe_allow_html=True
                )

        row_b = st.columns(2)
        with row_b[0]:
            college = st.text_input("College / University", placeholder="Your University", key=f"col_{idx}")
            if college.strip() and any(ch.isdigit() for ch in college):
                st.markdown(
                    "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ College name should contain letters only, no numbers.</p>",
                    unsafe_allow_html=True
                )
        with row_b[1]:
            grade = st.text_input("Grade / CGPA", placeholder="8.5 / 10", key=f"gr_{idx}")
            if grade.strip() and any(ch.isalpha() for ch in grade):
                st.markdown(
                    "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ Grade/CGPA should contain numbers only.</p>",
                    unsafe_allow_html=True
                )

        education.append({"degree": degree, "college": college, "year": year, "grade": grade})

st.markdown("<hr style='border-color:rgba(139,92,246,0.14)'>", unsafe_allow_html=True)

# ── Skills ─────────────────────────────────────────────────
st.markdown("### 🛠️ Skills")
sk1, sk2 = st.columns(2)
with sk1:
    skills = st.text_input("Technical Skills *", placeholder="Python, Java, HTML, CSS, SQL")
with sk2:
    softskills = st.text_input("Soft Skills", placeholder="Communication, Teamwork, Problem Solving")

st.markdown("<hr style='border-color:rgba(139,92,246,0.14)'>", unsafe_allow_html=True)

# ── Projects ───────────────────────────────────────────────
st.markdown("### 🚀 Projects")
st.markdown("<p style='color:#6B7280;font-size:0.85rem;'>Put each point on a new line in Description.</p>",
            unsafe_allow_html=True)
projects = []
for idx in range(3):
    label = f"Project {idx+1}" + (" *" if idx == 0 else " (optional)")
    with st.expander(f"💡 {label}", expanded=(idx == 0)):
        proj_name = st.text_input("Project Name", placeholder="AI Resume Analyzer",       key=f"pname_{idx}")
        proj_desc = st.text_area("Description",
                                  placeholder="Built a web app using Python and Streamlit\nImplemented ML model for job prediction",
                                  height=80, key=f"pdesc_{idx}")
        _desc_lines = [ln for ln in proj_desc.split("\n") if ln.strip()]
        if proj_desc.strip() and len(_desc_lines) < 2:
            st.markdown(
                "<p style='color:#DC2626;font-size:0.78rem;margin-top:-0.6rem;'>⚠️ Add more than one line (one point per line).</p>",
                unsafe_allow_html=True
            )
        proj_tech = st.text_input("Tech Stack",   placeholder="Python, Streamlit, SQLite", key=f"ptech_{idx}")
        projects.append({"name": proj_name, "description": proj_desc, "tech": proj_tech})

st.markdown("<hr style='border-color:rgba(139,92,246,0.14)'>", unsafe_allow_html=True)

# ── Certifications ─────────────────────────────────────────
st.markdown("### 🏆 Certifications & Achievements (optional)")
certifications = st.text_area("One per line",
                               placeholder="Python for Everybody — Coursera (2024)\nMachine Learning — Andrew Ng (2024)",
                               height=80)

st.markdown("<hr style='border-color:rgba(139,92,246,0.14)'>", unsafe_allow_html=True)

# ── Generate ───────────────────────────────────────────────
st.markdown("### ⚡ Generate Resume")

if st.button("🛠️ Generate My Resume", use_container_width=True, type="primary"):
    errors = []
    if not name.strip():            errors.append("Full Name is required.")
    elif any(ch.isdigit() for ch in name):
        errors.append("Name should not contain numbers.")
    if not email.strip():
        errors.append("Email is required.")
    elif not is_valid_email(email.strip()):
        errors.append("Please enter a valid email address (e.g. you@example.com).")
    if phone.strip():
        if not phone.strip().isdigit():
            errors.append("Phone number can only contain digits.")
        elif len(phone.strip()) != 10:
            errors.append("Phone number must be exactly 10 digits.")
    if linkedin.strip() and not is_valid_linkedin_url(linkedin.strip()):
        errors.append("Please enter a valid LinkedIn URL (e.g. linkedin.com/in/yourname).")
    if github.strip() and not is_valid_github_url(github.strip()):
        errors.append("Please enter a valid GitHub URL (e.g. github.com/yourname).")
    if occupation == "Select Occupation":  errors.append("Please select your current occupation.")
    if department == "Select Department":  errors.append("Please select your department.")
    if experience == "Select Experience":  errors.append("Please select your experience level.")
    if not summary.strip():         errors.append("Professional Summary is required.")
    elif len(_summary_sentences) < 2:
        errors.append("Professional Summary should have more than one sentence.")
    if not skills.strip():          errors.append("Technical Skills are required.")
    for e in education:
        if e["degree"].strip() and any(ch.isdigit() for ch in e["degree"]):
            errors.append("Degree should contain letters only, no numbers.")
        if e["college"].strip() and any(ch.isdigit() for ch in e["college"]):
            errors.append("College name should contain letters only, no numbers.")
        if e["year"].strip() and not _re.match(r"^\d{4}\s*-\s*\d{4}$", e["year"].strip()):
            errors.append("Education Year should be in the format 2020 - 2025.")
        if e["grade"].strip() and any(ch.isalpha() for ch in e["grade"]):
            errors.append("Education Grade/CGPA should contain numbers only.")

    if not projects[0]["name"].strip(): errors.append("At least 1 Project is required.")
    elif len([ln for ln in projects[0]["description"].split("\n") if ln.strip()]) < 2:
        errors.append("Project 1 Description should have more than one line.")

    if errors:
        for e in errors:
            st.error(e)
    else:
        with st.spinner("⚙️ Building your resume..."):
            try:
                data = {
                    "name":           name.strip(),
                    "email":          email.strip(),
                    "phone":          phone.strip(),
                    "linkedin":       linkedin.strip(),
                    "github":         github.strip(),
                    "occupation":     occupation.strip(),
                    "department":     department.strip(),
                    "experience":     experience.strip(),
                    "summary":        summary.strip(),
                    "education":      [e for e in education if e["degree"].strip()],
                    "skills":         skills.strip(),
                    "softskills":     softskills.strip(),
                    "projects":       [p for p in projects if p["name"].strip()],
                    "certifications": certifications.strip(),
                    "color":          chosen_color,
                }
                docx_bytes = generate_resume_docx(data)
                st.success("✅ Resume generated successfully!")
                st.download_button(
                    label="⬇️ Download Your Resume (.docx)",
                    data=docx_bytes,
                    file_name=f"resume_{name.strip().replace(' ','_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                st.info("💡 Open in Microsoft Word or Google Docs. Export as PDF before uploading to job portals.")
            except Exception as e:
                st.error(f"❌ Error: {e}")